"""
文档处理器
负责加载和预处理各种格式的文档
"""

import os
from pathlib import Path
from typing import List, Optional, Dict, Any
from dataclasses import dataclass, field

from src.utils import logger
from src.models.enums import KnowledgeCategory


@dataclass
class Document:
    """文档数据结构"""
    content: str                              # 文档内容
    source: str                               # 来源文件路径
    doc_id: str                               # 文档ID
    title: Optional[str] = None               # 文档标题
    category: Optional[KnowledgeCategory] = None  # 知识分类
    metadata: Dict[str, Any] = field(default_factory=dict)  # 元数据


class DocumentProcessor:
    """文档处理器 - 加载和预处理文档"""

    SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}

    def __init__(self, knowledge_base_path: str = "knowledge_base"):
        """
        初始化文档处理器

        Args:
            knowledge_base_path: 知识库根目录
        """
        self.knowledge_base_path = Path(knowledge_base_path)
        self.category_mapping = {
            "standards": KnowledgeCategory.STANDARDS,
            "formulas": KnowledgeCategory.FORMULAS,
            "operations": KnowledgeCategory.OPERATIONS,
            "cases": KnowledgeCategory.CASES,
            "faq": KnowledgeCategory.FAQ
        }

    def load_all_documents(self) -> List[Document]:
        """
        加载知识库中的所有文档

        Returns:
            文档列表
        """
        documents = []

        if not self.knowledge_base_path.exists():
            logger.warning(f"知识库目录不存在: {self.knowledge_base_path}")
            return documents

        for category_dir in self.knowledge_base_path.iterdir():
            if category_dir.is_dir():
                category_name = category_dir.name
                category = self.category_mapping.get(category_name)

                for file_path in category_dir.rglob("*"):
                    if file_path.is_file() and file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                        try:
                            doc = self._load_document(file_path, category)
                            if doc:
                                documents.append(doc)
                                logger.info(f"已加载文档: {file_path.name}")
                        except Exception as e:
                            logger.error(f"加载文档失败 {file_path}: {e}")

        logger.info(f"共加载 {len(documents)} 个文档")
        return documents

    def load_document(self, file_path: str) -> Optional[Document]:
        """
        加载单个文档

        Args:
            file_path: 文件路径

        Returns:
            Document对象
        """
        path = Path(file_path)
        if not path.exists():
            logger.error(f"文件不存在: {file_path}")
            return None

        # 尝试从路径推断分类
        category = None
        for part in path.parts:
            if part in self.category_mapping:
                category = self.category_mapping[part]
                break

        return self._load_document(path, category)

    def _load_document(self, file_path: Path, category: Optional[KnowledgeCategory] = None) -> Optional[Document]:
        """
        内部方法：加载单个文档

        Args:
            file_path: 文件路径
            category: 知识分类

        Returns:
            Document对象
        """
        suffix = file_path.suffix.lower()

        if suffix == ".md":
            content = self._load_markdown(file_path)
        elif suffix == ".txt":
            content = self._load_text(file_path)
        elif suffix == ".pdf":
            content = self._load_pdf(file_path)
        elif suffix == ".docx":
            content = self._load_docx(file_path)
        else:
            logger.warning(f"不支持的文件格式: {suffix}")
            return None

        if not content or not content.strip():
            logger.warning(f"文档内容为空: {file_path}")
            return None

        # 提取标题
        title = self._extract_title(content, file_path)

        # 生成文档ID
        doc_id = self._generate_doc_id(file_path)

        return Document(
            content=content,
            source=str(file_path),
            doc_id=doc_id,
            title=title,
            category=category,
            metadata={
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "extension": suffix
            }
        )

    def _load_markdown(self, file_path: Path) -> str:
        """加载Markdown文件"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(file_path, "r", encoding="gbk") as f:
                return f.read()

    def _load_text(self, file_path: Path) -> str:
        """加载纯文本文件"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="gbk") as f:
                return f.read()

    def _load_pdf(self, file_path: Path) -> str:
        """加载PDF文件"""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(str(file_path))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            return "\n\n".join(text_parts)
        except ImportError:
            logger.error("PyPDF2未安装，无法处理PDF文件")
            return ""
        except Exception as e:
            logger.error(f"PDF解析失败: {e}")
            return ""

    def _load_docx(self, file_path: Path) -> str:
        """加载Word文档"""
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(str(file_path))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except ImportError:
            logger.error("python-docx未安装，无法处理Word文件")
            return ""
        except Exception as e:
            logger.error(f"Word文档解析失败: {e}")
            return ""

    def _extract_title(self, content: str, file_path: Path) -> str:
        """
        提取文档标题

        优先从Markdown一级标题提取，否则使用文件名
        """
        lines = content.split("\n")
        for line in lines[:10]:  # 只检查前10行
            line = line.strip()
            if line.startswith("# "):
                return line[2:].strip()

        # 使用文件名（去掉扩展名）
        return file_path.stem

    def _generate_doc_id(self, file_path: Path) -> str:
        """生成文档ID"""
        import hashlib
        path_str = str(file_path.absolute())
        return hashlib.md5(path_str.encode()).hexdigest()[:12]


def create_document_processor(knowledge_base_path: str = "knowledge_base") -> DocumentProcessor:
    """创建文档处理器"""
    return DocumentProcessor(knowledge_base_path)
