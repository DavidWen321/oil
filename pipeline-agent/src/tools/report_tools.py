"""报告文件生成工具。"""

from __future__ import annotations

import os
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from src.utils import logger

BASE_DIR = Path(__file__).resolve().parents[2]
REPORT_DIR = BASE_DIR / "reports"
REPORT_DIR.mkdir(parents=True, exist_ok=True)


def generate_docx_report(report_data: Dict[str, Any]) -> Dict[str, Any]:
    """根据结构化报告数据生成 DOCX 文件。"""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)

    title = report_data.get("title", "管道运行分析报告")
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generate_time = report_data.get("generate_time") or datetime.now().isoformat()
    time_para = doc.add_paragraph(f"生成时间：{generate_time}")
    time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    summary = report_data.get("summary", "")
    if summary:
        doc.add_heading("摘要", level=1)
        doc.add_paragraph(str(summary))

    sections = report_data.get("sections", [])
    for section in sections:
        if not isinstance(section, dict):
            continue

        section_title = section.get("title", "章节")
        doc.add_heading(str(section_title), level=1)

        content = section.get("content", "")
        if content:
            for paragraph_text in str(content).split("\n"):
                paragraph_text = paragraph_text.strip()
                if paragraph_text:
                    doc.add_paragraph(paragraph_text)

        tables = section.get("tables", [])
        for table_data in tables:
            if not isinstance(table_data, dict):
                continue

            headers = table_data.get("headers", [])
            rows = table_data.get("rows", [])
            if not headers or not rows:
                continue

            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for idx, header in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.text = str(header)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            for row_idx, row in enumerate(rows):
                for col_idx, value in enumerate(row):
                    if col_idx < len(headers):
                        table.rows[row_idx + 1].cells[col_idx].text = str(value)

            doc.add_paragraph("")

        alerts = section.get("alerts", [])
        for alert in alerts:
            if not isinstance(alert, dict):
                continue
            msg = alert.get("message", "")
            level = alert.get("level", "info")
            prefix = {"warning": "⚠ 警告", "error": "❌ 严重", "info": "ℹ 提示"}.get(level, "ℹ 提示")
            doc.add_paragraph(f"【{prefix}】{msg}")

    recommendations = report_data.get("recommendations", [])
    if recommendations:
        doc.add_heading("优化建议", level=1)
        for idx, rec in enumerate(recommendations, start=1):
            doc.add_paragraph(f"{idx}. {rec}")

    file_name = f"report_{int(time.time())}.docx"
    file_path = REPORT_DIR / file_name
    doc.save(str(file_path))

    file_size = os.path.getsize(file_path)
    logger.info(f"报告已生成: {file_path}, 大小: {file_size} bytes")

    return {
        "file_path": str(file_path),
        "file_name": file_name,
        "file_size": file_size,
    }

