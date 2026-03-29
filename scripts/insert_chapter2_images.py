from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


DOC_PATH = Path(r"D:\oil\docs\智能管道能耗分析系统-使用说明书.docx")
DEFAULT_OUTPUT_PATH = Path(r"D:\oil\docs\智能管道能耗分析系统-使用说明书-第二章图片修正版.docx")
IMG_DIR = Path(r"D:\oil\docs\images")
PNG_DIR = IMG_DIR / "png"
CAPTION_TO_IMAGE = {
    "图1 前端功能设计图（图片预留）": ("图1 前端功能设计图", PNG_DIR / "图1-前端功能设计图.png"),
    "图1 前端功能设计图": ("图1 前端功能设计图", PNG_DIR / "图1-前端功能设计图.png"),
    "图2 后端功能设计图（图片预留）": ("图2 后端功能设计图", PNG_DIR / "图2-后端功能设计图.png"),
    "图2 后端功能设计图": ("图2 后端功能设计图", PNG_DIR / "图2-后端功能设计图.png"),
    "图3 系统架构图（图片预留）": ("图3 系统架构图", PNG_DIR / "图3-系统架构图.png"),
    "图3 系统架构图": ("图3 系统架构图", PNG_DIR / "图3-系统架构图.png"),
    "图4 系统整体设计图（图片预留）": ("图4 系统整体设计图", PNG_DIR / "图4-系统整体设计图.png"),
    "图4 系统整体设计图": ("图4 系统整体设计图", PNG_DIR / "图4-系统整体设计图.png"),
}


def convert_svgs() -> dict[str, Path]:
    result: dict[str, Path] = {}
    for original_caption, (_, png_path) in CAPTION_TO_IMAGE.items():
        if not png_path.exists():
            raise FileNotFoundError(png_path)
        result[original_caption] = png_path
    return result


def paragraph_has_drawing(paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing"))


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def insert_paragraph_after(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def set_caption_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def main() -> int:
    input_path = DOC_PATH
    output_path = Path(__import__("sys").argv[1]) if len(__import__("sys").argv) > 1 else input_path
    if not input_path.exists():
        raise FileNotFoundError(input_path)

    png_map = convert_svgs()
    doc = Document(str(input_path))

    for paragraph in list(doc.paragraphs):
        original_text = paragraph.text.strip()
        if original_text not in CAPTION_TO_IMAGE:
            continue

        new_caption, _ = CAPTION_TO_IMAGE[original_text]
        png_path = png_map[original_text]
        set_caption_text(paragraph, new_caption)

        next_element = paragraph._p.getnext()
        while next_element is not None:
            next_paragraph = Paragraph(next_element, paragraph._parent)
            if paragraph_has_drawing(next_paragraph):
                to_remove = next_paragraph
                next_element = next_element.getnext()
                remove_paragraph(to_remove)
                continue
            break

        image_para = insert_paragraph_after(paragraph)
        image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_para.add_run().add_picture(str(png_path), width=Inches(6.3))

    try:
        doc.save(str(output_path))
    except PermissionError:
        doc.save(str(DEFAULT_OUTPUT_PATH))
        output_path = DEFAULT_OUTPUT_PATH
    print(output_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
